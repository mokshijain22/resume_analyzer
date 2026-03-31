"""
resume_parser.py — Extract text from PDF and DOCX files.
Renamed from parser.py to avoid conflict with Python's built-in 'parser' module.
"""
import os


def extract_text(file_path: str) -> str:
    """
    Extract plain text from a PDF or DOCX file.
    Returns empty string on failure (caller checks for this).
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return _extract_pdf(file_path)
    elif ext == ".docx":
        return _extract_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _extract_pdf(path: str) -> str:
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        return "\n".join(text_parts)
    except Exception as e:
        raise RuntimeError(f"PDF extraction failed: {e}")


def _extract_docx(path: str) -> str:
    try:
        from docx import Document
        doc = Document(path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return "\n".join(parts)
    except Exception as e:
        raise RuntimeError(f"DOCX extraction failed: {e}")


if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        result = extract_text(sys.argv[1])
        print(f"Extracted {len(result)} characters")
        print(result[:500])